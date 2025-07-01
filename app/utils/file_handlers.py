"""
File handling utilities for processing uploaded documents
"""
import io
from pathlib import Path

import docx
import pypdf
from starlette.datastructures import UploadFile


async def extract_file_content(file: UploadFile) -> str:
    """
    Extract text content from uploaded file
    
    Args:
        file: Uploaded file object
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is not supported
        Exception: If there's an error processing the file
    """
    if not file or not file.filename:
        raise ValueError("No file provided")

    # Get file extension
    file_ext = Path(file.filename).suffix.lower()

    # Read file content into memory
    content_bytes = await file.read()

    # Reset file pointer for potential future reads
    await file.seek(0)

    try:
        if file_ext == '.txt':
            # Handle text files
            try:
                # Try UTF-8 first
                content = content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                # Fall back to latin-1 if UTF-8 fails
                content = content_bytes.decode('latin-1')

        elif file_ext == '.pdf':
            # Handle PDF files
            content = extract_pdf_content(content_bytes)

        elif file_ext == '.docx':
            # Handle Word documents
            content = extract_docx_content(content_bytes)

        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        return content.strip()

    except Exception as e:
        raise Exception(f"Error processing {file_ext} file: {e!s}")


def extract_pdf_content(content_bytes: bytes) -> str:
    """
    Extract text from PDF content
    
    Args:
        content_bytes: PDF file content as bytes
        
    Returns:
        Extracted text
    """
    try:
        pdf_file = io.BytesIO(content_bytes)
        pdf_reader = pypdf.PdfReader(pdf_file)

        # Check if PDF is encrypted
        if pdf_reader.is_encrypted:
            raise ValueError("Cannot extract text from encrypted PDF files")

        # Check if PDF has pages
        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF file has no pages")

        text_content = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)

        # Join and clean up text
        full_text = '\n'.join(text_content)

        # Check if we extracted any text
        if not full_text.strip():
            raise ValueError("No text content found in PDF. The PDF might contain only images.")

        return full_text

    except pypdf.errors.PdfReadError as e:
        raise ValueError(f"Invalid or corrupted PDF file: {e!s}")
    except Exception as e:
        raise Exception(f"Error extracting PDF content: {e!s}")


def extract_docx_content(content_bytes: bytes) -> str:
    """
    Extract text from Word document content
    
    Args:
        content_bytes: DOCX file content as bytes
        
    Returns:
        Extracted text
    """
    try:
        docx_file = io.BytesIO(content_bytes)
        doc = docx.Document(docx_file)

        text_content = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text:  # Avoid duplicates
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(' | '.join(row_text))

        # Check if we extracted any text
        if not text_content:
            raise ValueError("No text content found in Word document")

        return '\n'.join(text_content)

    except ValueError:
        raise
    except Exception as e:
        # Check for common issues
        if "BadZipFile" in str(e):
            raise ValueError("Invalid Word document format. The file might be corrupted or not a valid DOCX file.")
        raise Exception(f"Error extracting DOCX content: {e!s}")


def validate_file_size(file: UploadFile, max_size_mb: int = 10) -> bool:
    """
    Validate that uploaded file doesn't exceed size limit
    
    Args:
        file: Uploaded file object
        max_size_mb: Maximum allowed file size in megabytes
        
    Returns:
        True if file size is valid, False otherwise
    """
    if not file or not hasattr(file, 'file'):
        return False

    # Get file size
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()
    file.file.seek(0)  # Reset to beginning

    max_size_bytes = max_size_mb * 1024 * 1024
    return file_size <= max_size_bytes


def get_safe_filename(filename: str) -> str:
    """
    Generate a safe filename by removing special characters
    
    Args:
        filename: Original filename
        
    Returns:
        Safe filename
    """
    if not filename:
        return "unnamed_file"

    # Get base name and extension
    base_name = Path(filename).stem
    extension = Path(filename).suffix

    # Remove special characters from base name
    safe_base = "".join(c for c in base_name if c.isalnum() or c in ('-', '_'))

    # Ensure the name isn't empty
    if not safe_base:
        safe_base = "file"

    return f"{safe_base}{extension}"
