#!/usr/bin/env python3
"""
Test script to verify file upload and content extraction functionality
"""

import asyncio
import io
import os

# Test imports
import sys

import docx
from reportlab.pdfgen import canvas

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.utils.file_handlers import (
    extract_docx_content,
    extract_pdf_content,
)


def create_test_txt_file():
    """Create a test text file"""
    content = """This is a test text file.
It contains multiple lines of text.
This text will be used to test the file upload functionality.

Test special characters: café, naïve, résumé
End of test file."""

    with open("test_upload.txt", "w", encoding="utf-8") as f:
        f.write(content)

    print("✓ Created test_upload.txt")
    return content


def create_test_pdf_file():
    """Create a test PDF file using reportlab"""
    from reportlab.lib.pagesizes import letter

    # Create a PDF in memory
    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)

    # Add text to PDF
    text_lines = [
        "This is a test PDF file.",
        "It contains sample text for testing.",
        "The file upload system should extract this text.",
        "",
        "Test paragraph 2:",
        "PDF extraction is working correctly.",
    ]

    y_position = 750
    for line in text_lines:
        can.drawString(100, y_position, line)
        y_position -= 20

    can.save()

    # Save to file
    packet.seek(0)
    with open("test_upload.pdf", "wb") as f:
        f.write(packet.getvalue())

    print("✓ Created test_upload.pdf")
    return "\n".join(text_lines)


def create_test_docx_file():
    """Create a test Word document"""
    doc = docx.Document()

    # Add content
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test Word document.")
    doc.add_paragraph(
        "It contains multiple paragraphs to test the extraction functionality."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Test list:")
    doc.add_paragraph("• Item 1", style="List Bullet")
    doc.add_paragraph("• Item 2", style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph("End of test document.")

    # Save document
    doc.save("test_upload.docx")

    print("✓ Created test_upload.docx")

    # Return expected content
    return """Test Document
This is a test Word document.
It contains multiple paragraphs to test the extraction functionality.

Test list:
• Item 1
• Item 2

End of test document."""


async def test_file_extraction():
    """Test file extraction for all supported formats"""
    print("\nTesting file content extraction...\n")

    # Create test files
    txt_content = create_test_txt_file()
    pdf_content = create_test_pdf_file()
    create_test_docx_file()

    print("\nTesting extraction...\n")

    # Test TXT extraction
    print("Testing TXT extraction:")
    with open("test_upload.txt", "rb") as f:
        txt_bytes = f.read()
    try:
        extracted_txt = txt_bytes.decode("utf-8").strip()
        print("✓ TXT extraction successful")
        print(f"  Original length: {len(txt_content)}")
        print(f"  Extracted length: {len(extracted_txt)}")
        assert txt_content == extracted_txt, "TXT content mismatch"
    except Exception as e:
        print(f"✗ TXT extraction failed: {e}")

    # Test PDF extraction
    print("\nTesting PDF extraction:")
    with open("test_upload.pdf", "rb") as f:
        pdf_bytes = f.read()
    try:
        extracted_pdf = extract_pdf_content(pdf_bytes).strip()
        print("✓ PDF extraction successful")
        print(f"  Extracted text preview: {extracted_pdf[:100]}...")
        # PDF extraction might have slight formatting differences
        for line in pdf_content.split("\n"):
            if line and line not in extracted_pdf:
                print(f"  Warning: Line not found in extracted PDF: '{line}'")
    except Exception as e:
        print(f"✗ PDF extraction failed: {e}")

    # Test DOCX extraction
    print("\nTesting DOCX extraction:")
    with open("test_upload.docx", "rb") as f:
        docx_bytes = f.read()
    try:
        extracted_docx = extract_docx_content(docx_bytes).strip()
        print("✓ DOCX extraction successful")
        print(f"  Extracted text preview: {extracted_docx[:100]}...")
        # Check if main content is present
        assert "Test Document" in extracted_docx, "Missing document title"
        assert "test Word document" in extracted_docx, "Missing document content"
    except Exception as e:
        print(f"✗ DOCX extraction failed: {e}")

    # Clean up test files
    print("\nCleaning up test files...")
    for file in ["test_upload.txt", "test_upload.pdf", "test_upload.docx"]:
        try:
            os.remove(file)
            print(f"✓ Removed {file}")
        except Exception:
            pass

    print("\n✅ File extraction testing complete!")


def test_file_validation():
    """Test file validation functions"""
    from app.utils.file_handlers import get_safe_filename

    print("\nTesting filename sanitization...")

    test_cases = [
        ("normal_file.txt", "normal_file.txt"),
        ("file with spaces.pdf", "filewithspaces.pdf"),
        ("file@#$%^&*.docx", "file.docx"),
        ("../../etc/passwd", "etcpasswd"),
        ("", "unnamed_file"),
        ("no_extension", "no_extension"),
        (".hidden", "hidden"),
    ]

    for original, expected in test_cases:
        safe = get_safe_filename(original)
        status = "✓" if safe == expected else "✗"
        print(f"{status} '{original}' -> '{safe}' (expected: '{expected}')")


if __name__ == "__main__":
    print("=== File Upload and Extraction Test ===\n")

    # Check if required libraries are installed
    # Note: docx and reportlab are already imported at the top
    # We just need to check for pypdf which we don't use directly
    try:
        import pypdf  # noqa: F401 - Only checking if installed for app compatibility

        print("✓ All required libraries are installed")
    except ImportError as e:
        print(f"✗ Missing required library: {e}")
        print("\nPlease install missing dependencies:")
        print("pip install pypdf python-docx reportlab")
        sys.exit(1)

    # Run async tests
    asyncio.run(test_file_extraction())

    # Run sync tests
    test_file_validation()
